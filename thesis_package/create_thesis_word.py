#!/usr/bin/env python3
"""
Clean thesis DOCX exporter.

Builds a Word document from the canonical markdown source in
`thesis-draft.md` using a real markdown parser so the output
does not leak raw fences, horizontal rules, or pipe tables.
"""

from __future__ import annotations

import argparse
import re
from datetime import datetime
from pathlib import Path
from typing import Iterable

import mistune
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BLACK = RGBColor(0, 0, 0)

GREEK_MAP = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "epsilon": "ε",
    "sigma": "σ",
}


REPO_ROOT = Path(__file__).resolve().parents[1]
PKG = Path(__file__).resolve().parent
DEFAULT_INPUT = PKG / "THESIS_GROUNDED_MANUSCRIPT.md"
DEFAULT_OUTPUT = PKG / "output" / "THESIS_GROUNDED.docx"


def setup_document_styles(doc: Document) -> None:
    """Configure thesis-friendly document styles."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.first_line_indent = Inches(0)

    heading1 = doc.styles["Heading 1"]
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.page_break_before = False

    heading2 = doc.styles["Heading 2"]
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(6)

    heading3 = doc.styles["Heading 3"]
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.font.italic = False

    heading4 = doc.styles["Heading 4"]
    heading4.font.name = "Times New Roman"
    heading4.font.size = Pt(12)
    heading4.font.bold = True
    heading4.font.italic = False
    heading4.font.color.rgb = BLACK
    heading4.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    heading4.paragraph_format.first_line_indent = Inches(0)
    heading4.paragraph_format.space_before = Pt(6)
    heading4.paragraph_format.space_after = Pt(3)

    for list_style in ("List Bullet", "List Number"):
        lst = doc.styles[list_style]
        lst.font.name = "Times New Roman"
        lst.font.size = Pt(12)
        lst.paragraph_format.left_indent = Inches(0.5)
        lst.paragraph_format.first_line_indent = Inches(0)
        lst.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        lst.paragraph_format.space_after = Pt(0)

    styles = doc.styles
    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = styles["Normal"]
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(10.5)
        code_style.paragraph_format.first_line_indent = Inches(0)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.25)
        code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

    if "Quote Block" not in styles:
        quote_style = styles.add_style("Quote Block", WD_STYLE_TYPE.PARAGRAPH)
        quote_style.base_style = styles["Normal"]
        quote_style.font.italic = True
        quote_style.font.color.rgb = BLACK
        quote_style.paragraph_format.first_line_indent = Inches(0)
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.25)
        quote_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    if "Equation" not in styles:
        eq_style = styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
        eq_style.base_style = styles["Normal"]
        eq_style.font.name = "Times New Roman"
        eq_style.font.size = Pt(12)
        eq_style.font.italic = True
        eq_style.font.color.rgb = BLACK
        eq_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_style.paragraph_format.first_line_indent = Inches(0)
        eq_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        eq_style.paragraph_format.space_before = Pt(6)
        eq_style.paragraph_format.space_after = Pt(6)

    if "Equation Label" not in styles:
        eq_label = styles.add_style("Equation Label", WD_STYLE_TYPE.PARAGRAPH)
        eq_label.base_style = styles["Normal"]
        eq_label.font.name = "Times New Roman"
        eq_label.font.size = Pt(12)
        eq_label.font.bold = True
        eq_label.font.color.rgb = BLACK
        eq_label.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        eq_label.paragraph_format.first_line_indent = Inches(0)
        eq_label.paragraph_format.space_before = Pt(6)
        eq_label.paragraph_format.space_after = Pt(0)

    if "Caption" not in styles:
        cap = styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.base_style = styles["Normal"]
        cap.font.name = "Times New Roman"
        cap.font.size = Pt(11)
        cap.font.italic = True
        cap.font.color.rgb = BLACK
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.first_line_indent = Inches(0)
        cap.paragraph_format.space_before = Pt(6)
        cap.paragraph_format.space_after = Pt(12)
        cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if "Bibliography" not in styles:
        bib = styles.add_style("Bibliography", WD_STYLE_TYPE.PARAGRAPH)
        bib.base_style = styles["Normal"]
        bib.font.name = "Times New Roman"
        bib.font.size = Pt(12)
        bib.font.color.rgb = BLACK
        bib.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        bib.paragraph_format.left_indent = Inches(0.5)
        bib.paragraph_format.first_line_indent = Inches(-0.5)
        bib.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        bib.paragraph_format.space_before = Pt(0)
        bib.paragraph_format.space_after = Pt(0)

    if "Table Note" not in styles:
        note = styles.add_style("Table Note", WD_STYLE_TYPE.PARAGRAPH)
        note.base_style = styles["Normal"]
        note.font.name = "Times New Roman"
        note.font.size = Pt(10)
        note.font.italic = True
        note.font.color.rgb = BLACK
        note.paragraph_format.first_line_indent = Inches(0)
        note.paragraph_format.space_before = Pt(3)
        note.paragraph_format.space_after = Pt(9)
        note.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_cover_page(doc: Document, date_text: str) -> None:
    """Add a simple thesis cover page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)

    lines = [
        ("YUAN ZE UNIVERSITY", 16, True),
        ("College of Management", 14, False),
        ("Department of Finance", 14, False),
        ("", 12, False),
        ("A Thesis Submitted in Partial Fulfillment", 12, False),
        ("of the Requirements for the Degree of", 12, False),
        ("Master of Science in Finance", 12, True),
        ("", 12, False),
        ("ENERGY AS A CONSTRAINT:", 18, True),
        ("Credibility, Pricing, and Settlement", 18, True),
        ("in Energy-Linked Digital Finance", 18, True),
        ("", 12, False),
        ("Christopher Ongko", 14, True),
        ("Student ID: 1133958", 12, False),
        ("", 12, False),
        ("Advisor: Dr. De-Rong Kong (孔德蓉)", 12, False),
        ("", 12, False),
        (date_text, 12, False),
    ]
    for text, size, bold in lines:
        run = p.add_run(text + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold


def add_toc_section(doc: Document) -> None:
    """Insert a printable table of contents page (PDF-safe)."""
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)


def _insert_page_number_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_page_numbers(doc: Document) -> None:
    """Centered page numbers; cover page (first page of section 0) unnumbered."""
    for section in doc.sections:
        section.different_first_page_header_footer = True
        first_footer = section.first_page_footer
        if first_footer.paragraphs:
            first_footer.paragraphs[0].clear()
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK
        _insert_page_number_field(paragraph)


def add_toc_entries(doc: Document, page_map: dict[str, int]) -> None:
    """Render TOC lines with dot leaders and page numbers."""
    from pdf_toc import TOC_SPECS

    for label, _pattern in TOC_SPECS:
        page = page_map.get(label)
        if page is None:
            continue
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_after = Pt(5)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        title_run = paragraph.add_run(label)
        title_run.font.name = "Times New Roman"
        title_run.font.size = Pt(12)
        title_run.font.color.rgb = BLACK
        title_run.add_tab()
        page_run = paragraph.add_run(str(page))
        page_run.font.name = "Times New Roman"
        page_run.font.size = Pt(12)
        page_run.font.color.rgb = BLACK


def _list_item_text(item: dict) -> str:
    parts: list[str] = []
    for child in item.get("children", []):
        if child.get("type") in {"block_text", "paragraph"}:
            parts.append(inline_text(child.get("children", [])))
        elif "children" in child:
            parts.append(inline_text(child["children"]))
    return " ".join(p.strip() for p in parts if p.strip())


def inline_text(tokens: Iterable[dict]) -> str:
    """Flatten inline markdown tokens into plain text."""
    parts: list[str] = []
    for token in tokens:
        token_type = token["type"]
        if token_type == "text":
            parts.append(token.get("raw", ""))
        elif token_type in {"softbreak", "linebreak"}:
            parts.append("\n")
        elif token_type == "codespan":
            parts.append(token.get("raw", ""))
        elif "children" in token:
            parts.append(inline_text(token["children"]))
    return "".join(parts)


def render_inline(
    paragraph,
    tokens: Iterable[dict],
    *,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
    code: bool = False,
    plain_links: bool = False,
) -> None:
    """Render inline markdown tokens into a paragraph."""
    for token in tokens:
        token_type = token["type"]
        if token_type == "text":
            run = paragraph.add_run(token.get("raw", ""))
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(10.5)
        elif token_type in {"softbreak", "linebreak"}:
            paragraph.add_run().add_break()
        elif token_type == "strong":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=True or bold,
                italic=italic,
                underline=underline,
                code=code,
                plain_links=plain_links,
            )
        elif token_type == "emphasis":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=bold,
                italic=True or italic,
                underline=underline,
                code=code,
                plain_links=plain_links,
            )
        elif token_type == "codespan":
            run = paragraph.add_run(token.get("raw", ""))
            run.bold = bold
            run.italic = italic
            run.underline = underline
            run.font.name = "Courier New"
            run.font.size = Pt(10.5)
        elif token_type == "link":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=bold,
                italic=italic,
                underline=not plain_links and True,
                code=code,
                plain_links=plain_links,
            )
        elif "children" in token:
            render_inline(
                paragraph,
                token["children"],
                bold=bold,
                italic=italic,
                underline=underline,
                code=code,
                plain_links=plain_links,
            )


def _finalize_paragraph_runs(paragraph, *, size: float = 12) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = BLACK
        if run.font.name in {None, "Calibri", "Calibri Light", "Arial"}:
            run.font.name = "Times New Roman"
        if run.font.size is None or run.font.size.pt < 8:
            run.font.size = Pt(size)


def _apply_paragraph_format(paragraph, role: str) -> None:
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    if role == "body":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0.5)
    elif role == "bibliography":
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.left_indent = Inches(0.5)
        pf.first_line_indent = Inches(-0.5)
    elif role in {"abstract", "meta"}:
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0)
    elif role == "quote":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0)
        pf.left_indent = Inches(0.5)
        pf.right_indent = Inches(0.25)
    elif role == "list":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0)
    elif role == "table_note":
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.first_line_indent = Inches(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif role == "caption":
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Inches(0)
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        pf.space_before = Pt(6)
        pf.space_after = Pt(12)


def _blacken_paragraph(paragraph) -> None:
    _finalize_paragraph_runs(paragraph)


def _paragraph_style_hint(text: str) -> str:
    if text.startswith(("Keywords:", "JEL Codes:")):
        return "meta"
    if re.match(r"^Table \d+(\.\d+)?[.:]\s", text) and len(text) < 120:
        return "caption"
    if re.match(r"^Figure \d", text) and len(text) < 120:
        return "caption"
    if text.startswith("The following blocks are exported") or text.startswith(
        "Table 5.4 lists all indexed"
    ):
        return "table_note"
    return "body"


def add_body_paragraph(
    doc: Document,
    token: dict,
    *,
    bibliography: bool = False,
    abstract_mode: bool = False,
) -> None:
    children = token.get("children", [])
    if len(children) == 1 and children[0].get("type") == "image":
        add_image(doc, children[0])
        return

    if len(children) == 1 and children[0].get("type") == "emphasis":
        text = inline_text(children[0].get("children", [])).strip()
        if re.match(r"^Figure \d", text):
            paragraph = doc.add_paragraph(style="Caption")
            _apply_paragraph_format(paragraph, "caption")
            paragraph.add_run(text)
            _blacken_paragraph(paragraph)
            return

    text = inline_text(children).strip()
    if not text:
        return

    hint = _paragraph_style_hint(text)
    if bibliography:
        paragraph = doc.add_paragraph(style="Bibliography")
        role = "bibliography"
    elif abstract_mode and hint == "meta":
        paragraph = doc.add_paragraph()
        role = "meta"
    elif abstract_mode:
        paragraph = doc.add_paragraph()
        role = "abstract"
    elif hint == "caption":
        paragraph = doc.add_paragraph(style="Caption")
        role = "caption"
        paragraph.paragraph_format.keep_with_next = True
    elif hint == "table_note":
        paragraph = doc.add_paragraph(style="Table Note")
        role = "table_note"
    else:
        paragraph = doc.add_paragraph()
        role = "meta" if hint == "meta" else "body"
    _apply_paragraph_format(paragraph, role)
    render_inline(
        paragraph,
        token.get("children", []),
        plain_links=bibliography,
    )
    _blacken_paragraph(paragraph)


def _set_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def _set_portrait(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width


def _style_table_cell(paragraph, *, compact: bool) -> None:
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    size = Pt(8) if compact else Pt(10)
    for run in paragraph.runs:
        run.font.size = size
        if run.font.name in {None, "Calibri", "Calibri Light"}:
            run.font.name = "Times New Roman"
        run.font.color.rgb = BLACK


def format_equation_line(line: str) -> str:
    text = line.strip()
    for name, symbol in GREEK_MAP.items():
        text = re.sub(rf"\b{name}\b", symbol, text)
    text = text.replace(" * ", " ")
    text = re.sub(r"_\{([^}]+)\}", r"₍\1₎", text)
    text = text.replace("_t", "ₜ")
    text = text.replace("log(", "log(")
    return text


def add_equation_block(doc: Document, raw: str) -> None:
    lines = [ln for ln in raw.strip().splitlines() if ln.strip()]
    if not lines:
        return
    label = None
    body: list[str] = []
    for line in lines:
        if re.match(r"^Equation\s+\d", line.strip()):
            label = line.strip()
        else:
            body.append(format_equation_line(line))
    if label:
        label_p = doc.add_paragraph(style="Equation Label")
        label_p.add_run(label)
    for expr in body:
        eq_p = doc.add_paragraph(style="Equation")
        eq_p.add_run(expr)


def add_code_block(doc: Document, token: dict) -> None:
    raw = token.get("raw", "").rstrip()
    if not raw:
        return
    paragraph = doc.add_paragraph(style="Code Block")
    lines = raw.splitlines()
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10.5)
        if idx < len(lines) - 1:
            run.add_break()


def add_block_quote(doc: Document, token: dict) -> None:
    for child in token.get("children", []):
        if child["type"] == "paragraph":
            paragraph = doc.add_paragraph(style="Quote Block")
            _apply_paragraph_format(paragraph, "quote")
            render_inline(paragraph, child.get("children", []))
            _blacken_paragraph(paragraph)


def add_markdown_table(doc: Document, token: dict) -> None:
    head = None
    body_rows = []
    for child in token.get("children", []):
        if child["type"] == "table_head":
            head = child
        elif child["type"] == "table_body":
            body_rows = child.get("children", [])

    if head is None:
        return

    header_cells = head.get("children", [])
    rows = 1 + len(body_rows)
    cols = len(header_cells)
    compact = cols >= 5
    use_landscape = cols >= 8

    restore_section = None
    if use_landscape:
        restore_section = doc.sections[-1]
        landscape = doc.add_section()
        landscape.top_margin = Inches(0.75)
        landscape.bottom_margin = Inches(0.75)
        landscape.left_margin = Inches(0.75)
        landscape.right_margin = Inches(0.75)
        _set_landscape(landscape)

    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = True

    for col_idx, cell_token in enumerate(header_cells):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        render_inline(paragraph, cell_token.get("children", []), bold=True)
        _style_table_cell(paragraph, compact=compact)
        _blacken_paragraph(paragraph)

    for row_idx, row_token in enumerate(body_rows, start=1):
        for col_idx, cell_token in enumerate(row_token.get("children", [])):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            render_inline(paragraph, cell_token.get("children", []))
            _style_table_cell(paragraph, compact=compact)
            _blacken_paragraph(paragraph)

    if compact and restore_section is not None:
        portrait = doc.add_section()
        portrait.top_margin = restore_section.top_margin
        portrait.bottom_margin = restore_section.bottom_margin
        portrait.left_margin = restore_section.left_margin
        portrait.right_margin = restore_section.right_margin
        _set_portrait(portrait)


def add_image(doc: Document, token: dict) -> None:
    """Render markdown image as centered figure + caption."""
    attrs = token.get("attrs", {})
    url = (attrs.get("url") or "").strip()
    alt = (attrs.get("alt") or "Figure").strip()
    if not url:
        return

    candidates = [
        Path(url),
        PKG / url,
        REPO_ROOT / url,
        PKG / "empirical_results" / "figures" / Path(url).name,
    ]
    img_path = next((p for p in candidates if p.is_file()), None)
    if img_path is None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0)
        run = paragraph.add_run(f"[Missing figure: {url}]")
        run.italic = True
        return

    fig_para = doc.add_paragraph()
    fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig_para.paragraph_format.first_line_indent = Inches(0)
    fig_para.paragraph_format.space_before = Pt(6)
    fig_para.paragraph_format.keep_with_next = True
    run = fig_para.add_run()
    run.add_picture(str(img_path), width=Inches(5.0))


def add_list(doc: Document, token: dict) -> None:
    ordered = token.get("attrs", {}).get("ordered", False)
    depth = token.get("attrs", {}).get("depth", 0)
    style_name = "List Number" if ordered else "List Bullet"

    for index, item in enumerate(token.get("children", []), start=1):
        paragraph = doc.add_paragraph(style=style_name)
        paragraph.paragraph_format.left_indent = Inches(0.5 + 0.25 * depth)
        _apply_paragraph_format(paragraph, "list")

        item_children = item.get("children", [])
        added_text = False
        for child in item_children:
            child_type = child["type"]
            if child_type in {"block_text", "paragraph"}:
                render_inline(paragraph, child.get("children", []))
                added_text = True
            elif child_type == "list":
                if not added_text:
                    paragraph.add_run("")
                add_list(doc, child)
        if not added_text:
            paragraph.add_run("")
        _blacken_paragraph(paragraph)


def markdown_heading_level(md_level: int) -> int:
    return max(1, min(4, md_level - 1))


def build_docx(
    input_path: Path,
    output_path: Path,
    date_text: str,
    *,
    include_cover: bool = True,
    chapter_only: bool = False,
    toc_page_map: dict[str, int] | None = None,
) -> None:
    markdown = input_path.read_text(encoding="utf-8")
    parser = mistune.create_markdown(renderer="ast", plugins=["table"])
    tokens = parser(markdown)

    doc = Document()
    setup_document_styles(doc)
    if include_cover:
        add_cover_page(doc, date_text)

    started = False
    in_references = False
    in_abstract = False
    chapter_count = 0
    skip_front_matter_toc = False

    for token in tokens:
        token_type = token["type"]

        if not started:
            if token_type == "heading":
                heading_text = inline_text(token.get("children", [])).strip()
                level = token.get("attrs", {}).get("level", 2)
                if chapter_only and level == 2 and heading_text.startswith("Chapter "):
                    started = True
                elif level == 2 and heading_text == "Abstract":
                    started = True
                else:
                    continue
            else:
                continue

        if token_type == "heading":
            level = token.get("attrs", {}).get("level", 2)
            heading_text = inline_text(token.get("children", [])).strip()

            if not chapter_only and level == 2 and heading_text == "Table of Contents":
                add_toc_section(doc)
                if toc_page_map:
                    add_toc_entries(doc, toc_page_map)
                skip_front_matter_toc = True
                continue

            if skip_front_matter_toc and level == 2:
                skip_front_matter_toc = False
                doc.add_page_break()

            if level == 2 and heading_text.startswith("Chapter "):
                chapter_count += 1
            elif level == 2 and heading_text == "References":
                doc.add_page_break()
                in_references = True
                in_abstract = False
            elif in_references and level == 2:
                in_references = False
            elif level == 2 and heading_text == "Abstract":
                in_abstract = True
            elif level == 2 and heading_text != "Abstract":
                in_abstract = False

            doc.add_heading(heading_text, level=markdown_heading_level(level))
            _blacken_paragraph(doc.paragraphs[-1])
            if level == 2 and heading_text == "Abstract":
                doc.paragraphs[-1].paragraph_format.page_break_before = True
            if level == 2 and heading_text.startswith("Chapter ") and chapter_count > 1:
                doc.paragraphs[-1].paragraph_format.page_break_before = True
            continue

        if skip_front_matter_toc:
            continue

        if token_type == "paragraph":
            add_body_paragraph(
                doc,
                token,
                bibliography=in_references,
                abstract_mode=in_abstract,
            )
        elif token_type == "block_code":
            raw = token.get("raw", "").rstrip()
            if raw.lstrip().startswith("Equation"):
                add_equation_block(doc, raw)
            else:
                add_code_block(doc, token)
        elif token_type == "block_quote":
            add_block_quote(doc, token)
        elif token_type == "table":
            add_markdown_table(doc, token)
        elif token_type == "list":
            add_list(doc, token)
        elif token_type == "image":
            add_image(doc, token)
        elif token_type in {"blank_line", "thematic_break"}:
            continue

    add_page_numbers(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def build_docx_polished(
    input_path: Path,
    output_path: Path,
    date_text: str,
    *,
    include_cover: bool = True,
    chapter_only: bool = False,
) -> None:
    """Two-pass build: measure page numbers, then render TOC with dot leaders."""
    if chapter_only:
        build_docx(
            input_path,
            output_path,
            date_text,
            include_cover=include_cover,
            chapter_only=True,
        )
        return

    from export_pdf import docx_to_pdf
    from pdf_toc import extract_toc_page_map

    pass1_docx = output_path.with_suffix("._pass1.docx")
    pass1_pdf = output_path.with_suffix("._pass1.pdf")
    build_docx(
        input_path,
        pass1_docx,
        date_text,
        include_cover=include_cover,
        chapter_only=False,
        toc_page_map=None,
    )
    page_map: dict[str, int] = {}
    if docx_to_pdf(pass1_docx, pass1_pdf):
        page_map = extract_toc_page_map(pass1_pdf)
    build_docx(
        input_path,
        output_path,
        date_text,
        include_cover=include_cover,
        chapter_only=False,
        toc_page_map=page_map or None,
    )
    pass1_docx.unlink(missing_ok=True)
    pass1_pdf.unlink(missing_ok=True)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a clean thesis DOCX from markdown.")
    parser.add_argument("--input", default=str(DEFAULT_INPUT), help="Markdown input path.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="DOCX output path.")
    parser.add_argument(
        "--date-text",
        default=datetime.now().strftime("%B %Y"),
        help="Cover-page date text.",
    )
    args = parser.parse_args()

    build_docx(Path(args.input), Path(args.output), args.date_text)
    print(f"Saved clean thesis document to {args.output}")


if __name__ == "__main__":
    main()
