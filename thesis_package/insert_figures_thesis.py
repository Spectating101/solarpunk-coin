from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

print("=" * 70)
print("INSERTING FIGURES INTO THESIS")
print("=" * 70)

# Load the complete thesis
doc = Document('thesis_package/COMPLETE_THESIS_FINAL.docx')

# Figure mapping: where each figure should go
figures_to_insert = [
    {
        'file': 'thesis_package/empirical_results/comprehensive_ceir_analysis.png',
        'caption': 'Figure 2.1: Comprehensive CEIR Analysis Across Regimes',
        'after_text': '2.6 Main Results',
        'width': 6.0
    },
    {
        'file': 'thesis_package/empirical_results/triple_experiment_analysis.png',
        'caption': 'Figure 2.2: Three Natural Experiments (China Ban, ETH Merge, Kazakhstan)',
        'after_text': 'Geographic Concentration Matters',
        'width': 6.0
    },
    {
        'file': 'thesis_package/empirical_results/bitcoin_ceir_corrected.png',
        'caption': 'Figure 2.3: Bias-Corrected CEIR Regression Results',
        'after_text': 'bias-corrected',
        'width': 5.5
    },
    {
        'file': 'thesis_package/empirical_results/residualized_ceir.png',
        'caption': 'Figure 2.4: Residualized CEIR (Controlling for Momentum and Attention)',
        'after_text': 'horse-race',
        'width': 5.5
    },
    {
        'file': 'thesis_package/empirical_results/pricing_convergence_plots.png',
        'caption': 'Figure 3.1: Binomial Tree vs Monte Carlo Convergence',
        'after_text': '3.4 Monte Carlo Validation',
        'width': 6.0
    },
    {
        'file': 'thesis_package/empirical_results/fixed_trading_strategy.png',
        'caption': 'Figure 3.2: Option Pricing Across Strike Prices',
        'after_text': 'Extended Structures',
        'width': 5.5
    },
]

print(f"\n✓ Loaded document: {len(doc.paragraphs)} paragraphs")
print(f"✓ Found {len(figures_to_insert)} figures to insert\n")

# Check which figures exist
for fig in figures_to_insert:
    exists = os.path.exists(fig['file'])
    status = "✓" if exists else "✗"
    print(f"{status} {os.path.basename(fig['file'])}")

# Insert figures
inserted_count = 0
for fig_info in figures_to_insert:
    if not os.path.exists(fig_info['file']):
        print(f"\n⚠ Skipping {fig_info['file']} (not found)")
        continue
    
    # Find insertion point
    for i, para in enumerate(doc.paragraphs):
        if fig_info['after_text'] in para.text:
            print(f"\n✓ Inserting {fig_info['caption']}")
            print(f"  at paragraph {i}: '{para.text[:60]}...'")
            
            # Insert after this paragraph
            # Add the figure
            insert_para = doc.paragraphs[i]._element
            parent = insert_para.getparent()
            
            # Add space
            new_para = parent.insert(parent.index(insert_para) + 1, doc.add_paragraph()._element)
            
            # Add figure paragraph
            fig_para = doc.add_paragraph()
            run = fig_para.add_run()
            run.add_picture(fig_info['file'], width=Inches(fig_info['width']))
            fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Move to correct position
            parent.insert(parent.index(insert_para) + 2, fig_para._element)
            
            # Add caption
            caption_para = doc.add_paragraph()
            caption_run = caption_para.add_run(fig_info['caption'])
            caption_run.italic = True
            caption_run.font.size = Pt(10)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            parent.insert(parent.index(insert_para) + 3, caption_para._element)
            
            inserted_count += 1
            break

print(f"\n" + "=" * 70)
print(f"✓ Inserted {inserted_count}/{len(figures_to_insert)} figures")
print("=" * 70)

# Save the new version
output_path = 'thesis_package/COMPLETE_THESIS_WITH_FIGURES.docx'
doc.save(output_path)

print(f"\n✅ SAVED: {output_path}")

# Get stats
from docx import Document as Doc2
final_doc = Doc2(output_path)
words = sum(len(p.text.split()) for p in final_doc.paragraphs)
print(f"\n📊 FINAL STATS:")
print(f"   Words: {words:,}")
print(f"   Paragraphs: {len(final_doc.paragraphs):,}")
print(f"   Tables: {len(final_doc.tables)}")
print(f"   Estimated pages: {words // 250}")

