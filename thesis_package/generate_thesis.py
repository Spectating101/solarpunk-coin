#!/usr/bin/env python3
"""
Complete Thesis Document Generator
Creates a fully formatted, submission-ready Word document from thesis-draft.md
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
import re
import os

def setup_document_styles(doc):
    """Configure proper Word styles for academic thesis"""
    
    # Configure Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_after = Pt(0)
    paragraph_format.first_line_indent = Inches(0.5)
    
    # Configure Heading 1
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.size = Pt(16)
    heading1.font.bold = True
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)
    heading1.paragraph_format.keep_with_next = True
    
    # Configure Heading 2
    heading2 = doc.styles['Heading 2']
    heading2.font.name = 'Times New Roman'
    heading2.font.size = Pt(14)
    heading2.font.bold = True
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(6)
    
    # Configure Heading 3
    heading3 = doc.styles['Heading 3']
    heading3.font.name = 'Times New Roman'
    heading3.font.size = Pt(12)
    heading3.font.bold = True
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)

def add_cover_page(doc):
    """Add professional cover page"""
    
    # University name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('YUAN ZE UNIVERSITY')
    run.font.size = Pt(16)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('College of Management')
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Finance')
    run.font.size = Pt(14)
    
    # Add space
    for _ in range(4):
        doc.add_paragraph()
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ENERGY-BACKED DERIVATIVES:\n')
    run.font.size = Pt(18)
    run.font.bold = True
    
    run = p.add_run('From Empirical Validation to a Credible\n')
    run.font.size = Pt(18)
    run.font.bold = True
    
    run = p.add_run('Pricing-and-Contract Framework')
    run.font.size = Pt(18)
    run.font.bold = True
    
    # Add space
    for _ in range(4):
        doc.add_paragraph()
    
    # Author
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Christopher Ongko')
    run.font.size = Pt(14)
    run.font.bold = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Student ID: 1133958')
    run.font.size = Pt(12)
    
    # Add space
    for _ in range(3):
        doc.add_paragraph()
    
    # Advisor
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Advisor: Dr. De-Rong Kong (孔德蓉)')
    run.font.size = Pt(12)
    
    # Add space
    for _ in range(2):
        doc.add_paragraph()
    
    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('April 2025')
    run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()

def parse_markdown_to_sections(md_file_path):
    """Parse thesis-draft.md into structured sections"""
    
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    sections = []
    current_section = {'title': '', 'level': 0, 'content': []}
    
    lines = content.split('\n')
    
    for line in lines:
        # Check for headings
        if line.startswith('# '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {
                'title': line.replace('# ', '').strip(),
                'level': 1,
                'content': []
            }
        elif line.startswith('## '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {
                'title': line.replace('## ', '').strip(),
                'level': 2,
                'content': []
            }
        elif line.startswith('### '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {
                'title': line.replace('### ', '').strip(),
                'level': 3,
                'content': []
            }
        elif line.startswith('#### '):
            if current_section['content']:
                sections.append(current_section)
            current_section = {
                'title': line.replace('#### ', '').strip(),
                'level': 4,
                'content': []
            }
        else:
            current_section['content'].append(line)
    
    if current_section['content']:
        sections.append(current_section)
    
    return sections

def add_table_from_data(doc, table_title, data_rows, column_headers=None):
    """Add a properly formatted table to the document"""
    
    # Add table title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(table_title)
    run.font.bold = True
    run.font.size = Pt(11)
    
    # Determine dimensions
    if column_headers:
        num_cols = len(column_headers)
        num_rows = len(data_rows) + 1
    else:
        num_cols = len(data_rows[0]) if data_rows else 1
        num_rows = len(data_rows)
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Add headers
    if column_headers:
        for i, header in enumerate(column_headers):
            cell = table.rows[0].cells[i]
            cell.text = str(header)
            # Bold header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        start_row = 1
    else:
        start_row = 0
    
    # Add data
    for row_idx, row_data in enumerate(data_rows):
        for col_idx, cell_value in enumerate(row_data):
            table.rows[start_row + row_idx].cells[col_idx].text = str(cell_value)
    
    # Add spacing after table
    doc.add_paragraph()
    
    return table

def create_summary_statistics_table(doc):
    """Table 2.1: Summary Statistics by Regime"""
    
    data = [
        ['Bitcoin Price ($)', '14,820', '18,340', '32,640', '21,150'],
        ['CEIR (raw ratio)', '30.0', '17.2', '29.2', '12.9'],
        ['log(CEIR)', '3.273', '0.483', '3.281', '0.436'],
        ['30-day Forward Return (%, ann.)', '—', '80.7', '—', '58.2'],
        ['Mining HHI (geographic)', '0.42', '0.09', '0.18', '0.06'],
        ['Weighted Electricity Cost ($/kWh)', '0.059', '0.008', '0.065', '0.014'],
        ['Observations (weekly)', '129', '—', '202', '—']
    ]
    
    headers = ['Variable', 'Pre-Ban Mean', 'Pre-Ban SD', 'Post-Ban Mean', 'Post-Ban SD']
    
    add_table_from_data(doc, 'Table 2.1: Summary Statistics by Regime', data, headers)
    
    # Add note
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.font.italic = True
    run.font.size = Pt(10)
    run = p.add_run('Pre-ban period: January 2018 – June 2021 (N=129 weeks). Post-ban period: July 2021 – April 2025 (N=202 weeks). HHI = Herfindahl-Hirschman Index of geographic mining concentration.')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)

def create_bias_corrected_regression_table(doc):
    """Table 2.2: Bias-Corrected CEIR Predictive Regressions"""
    
    data = [
        ['log(CEIR)', '−0.165***', '−0.206***', '−0.199***'],
        ['', '(0.044)', '(0.042)', '(0.043)'],
        ['[log(CEIR)]²', '—', '—', '0.029'],
        ['', '—', '—', '(0.031)'],
        ['Fear & Greed Index', '—', '0.003***', '0.003***'],
        ['', '—', '(0.001)', '(0.001)'],
        ['û (AR residual)', '−0.162', '−0.206**', '−0.197**'],
        ['', '(0.105)', '(0.093)', '(0.095)'],
        ['Observations', '124', '124', '124'],
        ['R²', '0.051', '0.167', '0.168'],
    ]
    
    headers = ['Variable', '(1) Basic', '(2) + Controls', '(3) + Nonlinear']
    
    add_table_from_data(doc, 'Table 2.2: Bias-Corrected CEIR Predicts Returns During Concentrated Mining (Pre-Ban, Weekly)', data, headers)
    
    # Add note
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.font.italic = True
    run.font.size = Pt(10)
    run = p.add_run('Heteroskedasticity-robust (HC1) standard errors in parentheses. Amihud-Hurvich (2004) augmented regression. Dependent variable: 30-day forward return (%). *** p<0.01, ** p<0.05, * p<0.10. Pre-ban sample: January 2018 – June 2021, N=124 weekly observations.')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)

def create_structural_break_table(doc):
    """Table 2.3: Structural Break at China Mining Ban"""
    
    data = [
        ['log(CEIR)', '−0.060**', '−0.080**', '—'],
        ['', '(0.028)', '(0.031)', '—'],
        ['p-value', '[0.033]', '[0.011]', '—'],
        ['Observations', '200', '200', '—'],
        ['R²', '0.038', '0.039', '—'],
        ['Chow F-statistic', '—', '—', '4.786***'],
        ['p-value', '—', '—', '[0.0009]'],
    ]
    
    headers = ['', 'Post-Ban Basic', 'Post-Ban + Controls', 'Chow Test']
    
    add_table_from_data(doc, 'Table 2.3: Structural Break at the China Mining Ban', data, headers)
    
    # Add note
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.font.italic = True
    run.font.size = Pt(10)
    run = p.add_run('HC1 standard errors. Amihud-Hurvich augmented specification. *** p<0.01, ** p<0.05. Post-ban sample: July 2021 – April 2025, N=200 weeks. Chow test compares pre-ban (N=124) and post-ban (N=200) coefficients on log(CEIR) from specification (2).')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)

def create_pricing_validation_table(doc):
    """Table 3.X: Cross-Location Pricing Validation"""
    
    data = [
        ['Taiwan (primary)', '189%', '0.01918', '0.01886', '−0.00219 (credit)'],
        ['Saudi Arabia', '172%', '0.01841', '0.01808', '−0.00212 (credit)'],
        ['Arizona, USA', '165%', '0.01877', '0.01813', '−0.00241 (credit)'],
        ['Brazil', '198%', '0.03702', '0.03389', '−0.00640 (credit)'],
        ['Germany', '45%', '0.00234', '0.00212', '−0.00034 (credit)'],
    ]
    
    headers = ['Location', 'σ (%)', 'Call ($/kWh)', 'Put ($/kWh)', 'Collar Net Cost']
    
    add_table_from_data(doc, 'Table 3.5: Cross-Location Pricing Validation (ATM Call and Collar, T = 0.25 years)', data, headers)
    
    # Add note
    p = doc.add_paragraph()
    run = p.add_run('Note: ')
    run.font.italic = True
    run.font.size = Pt(10)
    run = p.add_run('S₀ = K (ATM). Taiwan σ calibrated from NASA POWER irradiance data (2019–2024) using explicit filtered preprocessing: 4-day rolling mean plus 1% absolute-return trim, yielding σ = 189.5%. Collar: buy put at 0.9×K, sell call at 1.1×K. Negative net cost indicates a structural credit to the producer. Pricing shown via binomial tree (N=400), with Taiwan base-case binomial-vs-Monte Carlo divergence of 2.08% at 20,000 paths. Risk-free rate r = 2.5%.')
    run.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)

def generate_complete_thesis(input_md_path, output_docx_path):
    """Main function to generate complete thesis document"""
    
    print("Initializing document...")
    doc = Document()
    
    print("Setting up styles...")
    setup_document_styles(doc)
    
    print("Adding cover page...")
    add_cover_page(doc)
    
    print("Reading thesis content...")
    with open(input_md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Parse content into sections
    sections = content.split('\n## ')
    
    print(f"Processing {len(sections)} major sections...")
    
    for i, section in enumerate(sections):
        if not section.strip():
            continue
        
        lines = section.split('\n')
        
        # First line is section title
        if i == 0:
            # Skip the main title line
            if lines[0].startswith('# '):
                lines = lines[1:]
        
        section_title = lines[0] if lines else ""
        section_content = '\n'.join(lines[1:]) if len(lines) > 1 else ""
        
        # Handle special sections
        if 'Abstract' in section_title:
            print("  Adding Abstract...")
            doc.add_heading('Abstract', level=1)
            paragraphs = section_content.split('\n\n')
            for para in paragraphs:
                if para.strip() and not para.strip().startswith('**'):
                    p = doc.add_paragraph(para.strip())
                    p.paragraph_format.first_line_indent = Inches(0)
        
        elif 'Table of Contents' in section_title:
            print("  Adding Table of Contents placeholder...")
            doc.add_heading('Table of Contents', level=1)
            p = doc.add_paragraph('[Table of Contents will be auto-generated in Word via Insert > Table of Contents]')
            p.paragraph_format.first_line_indent = Inches(0)
            doc.add_page_break()
        
        elif 'Chapter' in section_title or section_title.startswith('1.') or section_title.startswith('2.') or section_title.startswith('3.') or section_title.startswith('4.') or section_title.startswith('5.'):
            print(f"  Processing: {section_title[:50]}...")
            doc.add_heading(section_title, level=1)
            
            # Process subsections
            subsections = section_content.split('\n### ')
            
            for subsection in subsections:
                if not subsection.strip():
                    continue
                
                sublines = subsection.split('\n')
                subtitle = sublines[0]
                subcontent = '\n'.join(sublines[1:])
                
                if subtitle.strip():
                    doc.add_heading(subtitle, level=2)
                
                # Add tables where appropriate
                if 'Summary Statistics' in subtitle or 'Table 2.1' in subcontent:
                    create_summary_statistics_table(doc)
                elif 'Bias-Corrected' in subtitle or 'Table 2.2' in subcontent:
                    create_bias_corrected_regression_table(doc)
                elif 'Structural Break' in subtitle or 'Table 2.3' in subcontent:
                    create_structural_break_table(doc)
                elif 'Cross-Location' in subtitle or 'Global Validation' in subtitle:
                    create_pricing_validation_table(doc)
                
                # Add paragraphs
                paragraphs = subcontent.split('\n\n')
                for para in paragraphs:
                    if para.strip() and not para.strip().startswith('>') and not para.strip().startswith('**Table') and not para.strip().startswith('**Figure'):
                        # Handle bold text
                        if '**' in para:
                            p = doc.add_paragraph()
                            parts = para.split('**')
                            for idx, part in enumerate(parts):
                                if idx % 2 == 0:
                                    p.add_run(part)
                                else:
                                    run = p.add_run(part)
                                    run.font.bold = True
                        else:
                            doc.add_paragraph(para.strip())
        
        elif 'References' in section_title:
            print("  Adding References...")
            doc.add_page_break()
            doc.add_heading('References', level=1)
            
            # Parse references
            refs = section_content.strip().split('\n\n')
            for ref in refs:
                if ref.strip():
                    p = doc.add_paragraph(ref.strip())
                    p.paragraph_format.first_line_indent = Inches(-0.5)
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(6)
    
    print(f"Saving document to {output_docx_path}...")
    doc.save(output_docx_path)
    print("✓ Complete thesis document generated successfully!")

if __name__ == '__main__':
    input_file = 'thesis-draft.md'
    output_file = 'thesis_package/COMPLETE_THESIS_SUBMISSION_READY.docx'
    
    generate_complete_thesis(input_file, output_file)
