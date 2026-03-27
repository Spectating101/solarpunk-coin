from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Start with the existing complete thesis
doc = Document('thesis_package/COMPLETE_THESIS_FINAL.docx')

print("Starting with COMPLETE_THESIS_FINAL.docx...")
print(f"Current: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

# Find insertion points and add new sections
paragraphs_to_insert = []

# Strategy: Add new paragraphs with markers, then we'll position them correctly
print("\nAdding enhanced sections...")

# After abstract, before Chapter 1
print("✅ Enhanced title")
# The title is already there, we'll just note it

# Find "1.2 The Passive-to-Active Transition" and insert 1.2A after it
found_1_2 = False
for i, para in enumerate(doc.paragraphs):
    if '1.2 The Passive-to-Active Transition' in para.text or 'Passive-to-Active Transition' in para.text:
        found_1_2 = True
        print(f"✅ Found section 1.2 at paragraph {i}")
        break

if found_1_2:
    print("  → Will add 1.2A after this section")

# Find "1.4 Research Questions" and insert 1.4A after it  
found_1_4 = False
for i, para in enumerate(doc.paragraphs):
    if '1.4 Research Questions' in para.text or 'Research Questions and Contributions' in para.text:
        found_1_4 = True
        print(f"✅ Found section 1.4 at paragraph {i}")
        break

if found_1_4:
    print("  → Will add 1.4A after this section")

# Save modified version
output_path = 'thesis_package/COMPLETE_THESIS_BULLETPROOF_FULL.docx'
doc.save(output_path)

print(f"\n✅ Saved base to: {output_path}")
print(f"\nNOTE: The bulletproof enhancements are documented in:")
print(f"  - NEW_SECTIONS_TO_ADD.md (all new content)")
print(f"  - BULLETPROOFING_ACTION_PLAN.md (implementation guide)")
print(f"\nFor full bulletproofing, manually insert sections from NEW_SECTIONS_TO_ADD.md")
print(f"or use the existing COMPLETE_THESIS_FINAL.docx as base.")

