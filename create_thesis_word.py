#!/usr/bin/env python3
"""
Clean thesis DOCX exporter.

Builds a Word document from the canonical markdown source in
`thesis_package/thesis-draft.md` using a real markdown parser so the output
does not leak raw fences, horizontal rules, or pipe tables.
"""

from __future__ import annotations

import argparse
from datetime import datetime
from pathlib import Path
from typing import Iterable

import mistune
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
DEFAULT_INPUT = ROOT / "thesis_package" / "thesis-draft.md"
DEFAULT_OUTPUT = ROOT / "thesis_package" / "THESIS_COMPLETE_FINAL.docx"


def setup_document_styles(doc: Document) -> None:
    """Configure thesis-friendly document styles."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    for style_name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        style.paragraph_format.first_line_indent = Inches(0)

    heading1 = doc.styles["Heading 1"]
    heading1.paragraph_format.space_before = Pt(24)
    heading1.paragraph_format.space_after = Pt(12)

    heading2 = doc.styles["Heading 2"]
    heading2.paragraph_format.space_before = Pt(18)
    heading2.paragraph_format.space_after = Pt(6)

    heading3 = doc.styles["Heading 3"]
    heading3.paragraph_format.space_before = Pt(12)
    heading3.paragraph_format.space_after = Pt(6)
    heading3.font.italic = True

    styles = doc.styles
    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = styles["Normal"]
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(10.5)
        code_style.paragraph_format.first_line_indent = Inches(0)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.right_indent = Inches(0.25)
        code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)

    if "Quote Block" not in styles:
        quote_style = styles.add_style("Quote Block", WD_STYLE_TYPE.PARAGRAPH)
        quote_style.base_style = styles["Normal"]
        quote_style.font.italic = True
        quote_style.paragraph_format.first_line_indent = Inches(0)
        quote_style.paragraph_format.left_indent = Inches(0.5)
        quote_style.paragraph_format.right_indent = Inches(0.25)


def add_cover_page(doc: Document, date_text: str) -> None:
    """Add a simple thesis cover page."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)

    lines = [
        ("YUAN ZE UNIVERSITY", 16, True),
        ("College of Management", 14, False),
        ("Department of Finance", 14, False),
        ("", 12, False),
        ("", 12, False),
        ("ENERGY-BACKED DERIVATIVES:", 18, True),
        ("From Empirical Validation to a Credible", 18, True),
        ("Pricing-and-Contract Framework", 18, True),
        ("", 12, False),
        ("", 12, False),
        ("Christopher Ongko", 14, True),
        ("Student ID: 1133958", 12, False),
        ("", 12, False),
        ("", 12, False),
        ("Advisor: Dr. De-Rong Kong (孔德蓉)", 12, False),
        ("", 12, False),
        (date_text, 12, False),
    ]
    for text, size, bold in lines:
        run = p.add_run(text + "\n")
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold

    doc.add_page_break()


def add_toc_section(doc: Document) -> None:
    """Insert a TOC field that Word can update."""
    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_text = OxmlElement("w:t")
    fld_text.text = "Update field in Word to populate the table of contents."

    fld_sep_run = OxmlElement("w:r")
    fld_sep_run.append(fld_text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = p.add_run()._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_sep_run)
    r.append(fld_end)

    note = doc.add_paragraph()
    note.paragraph_format.first_line_indent = Inches(0)
    run = note.add_run("In Word: right-click the table of contents and choose 'Update Field'.")
    run.font.italic = True
    run.font.size = Pt(10)

    doc.add_page_break()


def inline_text(tokens: Iterable[dict]) -> str:
    """Flatten inline markdown tokens into plain text."""
    parts: list[str] = []
    for token in tokens:
        token_type = token["type"]
        if token_type == "text":
            parts.append(token.get("raw", ""))
        elif token_type in {"softbreak", "linebreak"}:
            parts.append("\n")
        elif token_type == "codespan":
            parts.append(token.get("raw", ""))
        elif "children" in token:
            parts.append(inline_text(token["children"]))
    return "".join(parts)


def render_inline(paragraph, tokens: Iterable[dict], *, bold: bool = False,
                  italic: bool = False, underline: bool = False,
                  code: bool = False) -> None:
    """Render inline markdown tokens into a paragraph."""
    for token in tokens:
        token_type = token["type"]
        if token_type == "text":
            run = paragraph.add_run(token.get("raw", ""))
            run.bold = bold
            run.italic = italic
            run.underline = underline
            if code:
                run.font.name = "Courier New"
                run.font.size = Pt(10.5)
        elif token_type in {"softbreak", "linebreak"}:
            paragraph.add_run().add_break()
        elif token_type == "strong":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=True or bold,
                italic=italic,
                underline=underline,
                code=code,
            )
        elif token_type == "emphasis":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=bold,
                italic=True or italic,
                underline=underline,
                code=code,
            )
        elif token_type == "codespan":
            run = paragraph.add_run(token.get("raw", ""))
            run.bold = bold
            run.italic = italic
            run.underline = underline
            run.font.name = "Courier New"
            run.font.size = Pt(10.5)
        elif token_type == "link":
            render_inline(
                paragraph,
                token.get("children", []),
                bold=bold,
                italic=italic,
                underline=True,
                code=code,
            )
        elif "children" in token:
            render_inline(
                paragraph,
                token["children"],
                bold=bold,
                italic=italic,
                underline=underline,
                code=code,
            )


def add_body_paragraph(doc: Document, token: dict) -> None:
    text = inline_text(token.get("children", [])).strip()
    if not text:
        return

    paragraph = doc.add_paragraph()
    if text.startswith("Table ") or text.startswith("Figure "):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Inches(0)
    elif text.startswith("Note:") or text.startswith("Keywords:") or text.startswith("JEL Codes:"):
        paragraph.paragraph_format.first_line_indent = Inches(0)
    render_inline(paragraph, token.get("children", []))


def add_code_block(doc: Document, token: dict) -> None:
    raw = token.get("raw", "").rstrip()
    if not raw:
        return
    paragraph = doc.add_paragraph(style="Code Block")
    lines = raw.splitlines()
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(10.5)
        if idx < len(lines) - 1:
            run.add_break()


def add_block_quote(doc: Document, token: dict) -> None:
    for child in token.get("children", []):
        if child["type"] == "paragraph":
            paragraph = doc.add_paragraph(style="Quote Block")
            render_inline(paragraph, child.get("children", []))


def add_markdown_table(doc: Document, token: dict) -> None:
    head = None
    body_rows = []
    for child in token.get("children", []):
        if child["type"] == "table_head":
            head = child
        elif child["type"] == "table_body":
            body_rows = child.get("children", [])

    if head is None:
        return

    header_cells = head.get("children", [])
    rows = 1 + len(body_rows)
    cols = len(header_cells)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"

    for col_idx, cell_token in enumerate(header_cells):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        render_inline(paragraph, cell_token.get("children", []), bold=True)

    for row_idx, row_token in enumerate(body_rows, start=1):
        for col_idx, cell_token in enumerate(row_token.get("children", [])):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            render_inline(paragraph, cell_token.get("children", []))

    doc.add_paragraph("")


def add_list(doc: Document, token: dict) -> None:
    ordered = token.get("attrs", {}).get("ordered", False)
    depth = token.get("attrs", {}).get("depth", 0)
    style_name = "List Number" if ordered else "List Bullet"

    for index, item in enumerate(token.get("children", []), start=1):
        paragraph = doc.add_paragraph(style=style_name)
        paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
        paragraph.paragraph_format.first_line_indent = Inches(0)

        item_children = item.get("children", [])
        added_text = False
        for child in item_children:
            child_type = child["type"]
            if child_type in {"block_text", "paragraph"}:
                render_inline(paragraph, child.get("children", []))
                added_text = True
            elif child_type == "list":
                if not added_text:
                    paragraph.add_run("")
                add_list(doc, child)
        if not added_text:
            paragraph.add_run("")


def markdown_heading_level(md_level: int) -> int:
    return max(1, min(4, md_level - 1))


def build_docx(input_path: Path, output_path: Path, date_text: str) -> None:
    markdown = input_path.read_text(encoding="utf-8")
    parser = mistune.create_markdown(renderer="ast", plugins=["table"])
    tokens = parser(markdown)

    doc = Document()
    setup_document_styles(doc)
    add_cover_page(doc, date_text)

    started = False
    skip_toc = False
    chapter_count = 0

    for token in tokens:
        token_type = token["type"]

        if not started:
            if token_type == "heading":
                heading_text = inline_text(token.get("children", [])).strip()
                if token.get("attrs", {}).get("level") == 2 and heading_text == "Abstract":
                    started = True
                else:
                    continue
            else:
                continue

        if token_type == "heading":
            level = token.get("attrs", {}).get("level", 2)
            heading_text = inline_text(token.get("children", [])).strip()

            if skip_toc and not (level == 2 and heading_text != "Table of Contents"):
                continue
            if skip_toc and level == 2 and heading_text != "Table of Contents":
                skip_toc = False

            if level == 2 and heading_text == "Table of Contents":
                add_toc_section(doc)
                skip_toc = True
                continue

            if level == 2 and heading_text.startswith("Chapter "):
                chapter_count += 1
                if chapter_count > 1:
                    doc.add_page_break()
            elif level == 2 and heading_text == "References":
                doc.add_page_break()

            doc.add_heading(heading_text, level=markdown_heading_level(level))
            continue

        if skip_toc:
            continue

        if token_type == "paragraph":
            add_body_paragraph(doc, token)
        elif token_type == "block_code":
            add_code_block(doc, token)
        elif token_type == "block_quote":
            add_block_quote(doc, token)
        elif token_type == "table":
            add_markdown_table(doc, token)
        elif token_type == "list":
            add_list(doc, token)
        elif token_type in {"blank_line", "thematic_break"}:
            continue

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a clean thesis DOCX from markdown.")
    parser.add_argument("--input", default=str(DEFAULT_INPUT), help="Markdown input path.")
    parser.add_argument("--output", default=str(DEFAULT_OUTPUT), help="DOCX output path.")
    parser.add_argument(
        "--date-text",
        default=datetime.now().strftime("%B %Y"),
        help="Cover-page date text.",
    )
    args = parser.parse_args()

    build_docx(Path(args.input), Path(args.output), args.date_text)
    print(f"Saved clean thesis document to {args.output}")


if __name__ == "__main__":
    main()
